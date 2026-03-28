import streamlit as st
import base64
import os
import pdfplumber
import docx
import pypdf
import random
try:
    from rapidocr_onnxruntime import RapidOCR
    RAPIDOCR_AVAILABLE = True
except ImportError:
    RAPIDOCR_AVAILABLE = False

import numpy as np

def get_ocr_engine():
    """Lazy-load the OCR engine only when needed."""
    if not RAPIDOCR_AVAILABLE:
        return None
    try:
        return RapidOCR()
    except Exception:
        return None

from fpdf import FPDF
from PIL import Image
import re
import html
import string

st.set_page_config(page_title="Resume Analyzer", layout="centered")

# -------------------------------
# Global Styling
# -------------------------------
st.markdown("""
<style>
    .main-title {
        color: #9400D3;
        text-align: center;
        font-size: 40px;
        font-weight: bold;
        margin-bottom: 10px;
    }
    .subtitle {
        color: #34495e;
        text-align: center;
        font-size: 18px;
        margin-bottom: 20px;
    }
    .section-title {
        color: #ffffff;
        font-size: 24px;
        font-weight: bold;
        margin-top: 30px;
        margin-bottom: 15px;
        border-bottom: 2px solid rgba(255,255,255,0.3);
        padding-bottom: 5px;
    }
    .score-container {
        background: rgba(255, 255, 255, 0.1);
        backdrop-filter: blur(10px);
        padding: 20px;
        border-radius: 12px;
        text-align: center;
        font-size: 24px;
        font-weight: bold;
        color: #ffffff;
        margin-top: 20px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        border: 1px solid rgba(255, 255, 255, 0.2);
    }
    .text-container {
        height: 300px;
        overflow-y: scroll;
        border: 1px solid rgba(255, 255, 255, 0.2);
        border-radius: 8px;
        padding: 15px;
        background: rgba(0, 0, 0, 0.2);
        white-space: pre-wrap;
        box-shadow: inset 0 2px 4px rgba(0,0,0,0.1);
        color: #ffffff !important;
    }
    .stProgress > div > div > div > div {
        background-color: #27ae60;
    }
    .metric-card {
        background: rgba(255, 255, 255, 0.1);
        backdrop-filter: blur(10px);
        border-radius: 8px;
        padding: 15px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        border: 1px solid rgba(255, 255, 255, 0.2);
        text-align: center;
        margin-bottom: 10px;
    }
    .metric-card h4 {
        margin: 0;
        color: #e0e0e0 !important;
        font-size: 16px;
    }
    .metric-card h2 {
        margin: 10px 0 0 0;
        color: #ffffff !important;
        font-size: 28px;
    }
</style>
""", unsafe_allow_html=True)

# -------------------------------
# Extract Resume Text
# -------------------------------
def extract_text(file):
    text = ""
    error_log = []

    if file.type == "application/pdf":
        try:
            # Method 1: pdfplumber
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
            
            # Method 2: pypdf (fallback if pdfplumber failed)
            if not text.strip():
                file.seek(0)
                reader = pypdf.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # Method 3: OCR (fallback if traditional extraction failed)
            if not text.strip():
                file.seek(0)
                with pdfplumber.open(file) as pdf:
                    for page in pdf.pages:
                        try:
                            im = page.to_image(resolution=300).original
                            img_np = np.array(im)
                            ocr_engine = get_ocr_engine()
                            if ocr_engine:
                                result, _ = ocr_engine(img_np)
                                if result:
                                    text += "\n".join([line[1] for line in result]) + "\n"
                            else:
                                error_log.append("OCR Engine (RapidOCR) not available.")
                        except Exception as e:
                            error_log.append(f"OCR Error: {str(e)}")
        except Exception as e:
            error_log.append(f"PDF Analysis Error: {str(e)}")

    elif file.type.endswith("document.wordprocessingml.document"):
        try:
            doc = docx.Document(file)
            for p in doc.paragraphs:
                if p.text:
                    text += p.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + "\n"
        except Exception as e:
            error_log.append(f"Word Analysis Error: {str(e)}")

    elif file.type.startswith("image"):
        try:
            img = Image.open(file)
            img_np = np.array(img.convert('RGB'))
            ocr_engine = get_ocr_engine()
            if ocr_engine:
                result, _ = ocr_engine(img_np)
                if result:
                    text += "\n".join([line[1] for line in result]) + "\n"
            else:
                error_log.append("OCR Engine (RapidOCR) not available.")
        except Exception as e:
            error_log.append(f"Image Analysis Error: {str(e)}")
            
    if not text.strip() and error_log:
        for err in error_log:
            st.sidebar.error(err)
        st.error("🚨 **System Failure**: Could not extract any text from the uploaded file. Please ensure it is a valid, unencrypted document.")

    return text


# -------------------------------
# Detect Resume Sections
# -------------------------------
def analyze_resume(text):
    t = text.lower()
    has_email = bool(re.search(r'[\w\.-]+@[\w\.-]+', text))
    has_phone = bool(re.search(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]', text))
    has_linkedin = "linkedin.com" in t

    sections_found = {
        "Contact Info": has_email or has_phone or has_linkedin,
        "Skills": bool(re.search(r'\b(skills|technical|technologies|proficiencies|expertise|knowledge)\b', t)),
        "Experience": bool(re.search(r'\b(experience|work history|employment|career|professional background)\b', t)),
        "Education": bool(re.search(r'\b(education|academic|university|degree|qualifications)\b', t)),
        "Projects": bool(re.search(r'\b(projects|portfolio|personal work|accomplishments)\b', t)),
        "Summary": bool(re.search(r'\b(summary|objective|profile|about me|professional profile)\b', t))
    }
    return sections_found


# -------------------------------
# Resume Score Calculation
# -------------------------------
def calculate_score(sections, text, target_role=""):
    details = {}
    t = text.lower()
    tr = target_role.lower()

    formatting = 0
    if len(text.split("\n")) > 12: formatting += 10
    if any(char in text for char in ["•", "-", "*", "·", "▪", "–", "—"]): formatting += 10
    details["Formatting"] = formatting

    content = 0
    words = len(text.split())
    if words > 200: content += 10
    if sections.get("Contact Info", False): content += 10
    details["Content"] = content

    skills_score = 0
    if sections.get("Skills", False): skills_score += 10
    keywords = [
        "python", "java", "javascript", "react", "node", "sql", "nosql", "aws",
        "azure", "docker", "kubernetes", "machine learning", "data", "ai", "agile", "scrum", "git", "ci/cd", 
        "project management", "leadership", "communication", "problem solving", "analytics"
    ]
    # Add words from target role
    role_words = [w for w in tr.split() if len(w) > 3]
    keywords.extend(role_words)

    matches = [kw for kw in keywords if kw in t]
    skills_score += min(20, len(matches) * 2) 
    details["Skills"] = skills_score

    exp_score = 0
    if sections.get("Experience", False): exp_score += 15
    impact_verbs = ["managed", "led", "developed", "increased", "reduced", "spearheaded", "optimized", "implemented"]
    if any(v in t for v in impact_verbs): exp_score += 5
    if re.search(r'\d+%', text) or re.search(r'\$\d+', text) or re.search(r'\b(years|months)\b', t):
        exp_score += 10
    details["Experience"] = exp_score

    total = formatting + content + skills_score + exp_score
    return details, min(total, 100)


# -------------------------------
# Suggestions
# -------------------------------
def generate_suggestions(sections, text, target_role=""):
    suggestions = []

    if target_role:
        suggestions.append(f"Ensure your summary highlights your exact qualifications for the '{target_role}' role.")

    if not sections["Summary"]:
        suggestions.append("Add a professional summary at the top.")

    if not sections["Skills"]:
        suggestions.append("Include a technical skills section.")
    
    if target_role:
        tr_words = [w for w in target_role.lower().split() if len(w) > 3]
        missing_role_words = [w for w in tr_words if w not in text.lower()]
        if missing_role_words:
            suggestions.append(f"Consider integrating these keywords related to your target role: {', '.join(missing_role_words)}")

    if not sections["Projects"]:
        suggestions.append("Add project experience.")

    if not re.search(r'\d+', text):
        suggestions.append("Use measurable achievements (example: improved performance by 20%).")

    if len(text.split()) < 200:
        suggestions.append("Resume is short. Add more detailed experience.")

    return suggestions


# -------------------------------
# Extract Bullet Points
# -------------------------------
def extract_bullets(text):
    bullets = []
    bullet_chars = ("•", "-", "*", "·", "▪", "–", "—", ">", "o ", "e ", "■", "♦")

    for line in text.split("\n"):
        line = line.strip()
        word_count = len(line.split())
        
        if word_count >= 5:
            if line.startswith(bullet_chars):
                bullets.append(line)
            elif re.match(r'^[A-Z][a-z]+ed\b', line):
                bullets.append(line)
    
    return list(dict.fromkeys(bullets))[:15]


# -------------------------------
# Deep Analysis & Rewrite of Bullet Points
# -------------------------------
def improve_bullets(bullets, target_role="", jd_text=""):
    improved = []
    
    context_text = f"{target_role} {jd_text}".lower()
    context_text = re.sub(f'[{re.escape(string.punctuation)}]', ' ', context_text)
    words = context_text.split()
    stop_words = {"the", "and", "to", "of", "in", "for", "is", "with", "on", "as", "an", "at", "this", "that", "it", "or", "from", "be", "are", "you", "your", "we", "our", "will", "can"}
    role_keywords = list(set([w for w in words if w not in stop_words and len(w) > 3]))
    
    # We must give 4-5 bullet points
    # Select randomly if there's an abundance, otherwise 5
    if len(bullets) > 5:
        bullets = random.sample(bullets, 5)
    
    # If no real bullets found, fallback mechanism will trigger downstream, 
    # but we will loop through whatever we have
    for bullet in bullets:
        clean_bullet = re.sub(r'^[\s\•\-\*\·\▪\–\—\>]+', '', bullet).strip()
        if not clean_bullet: continue

        bullet_words = clean_bullet.split()
        first_word = bullet_words[0].lower() if bullet_words else ""
        has_metrics = bool(re.search(r'\d+', clean_bullet) or '%' in clean_bullet or '$' in clean_bullet)
        
        weak_to_strong = {
            "helped": "Spearheaded", "worked": "Engineered", "did": "Executed", "was": "Directed",
            "responsible": "Orchestrated", "assisted": "Accelerated", "contributed": "Catalyzed",
            "managed": "Optimized", "used": "Leveraged", "made": "Constructed", "led": "Pioneered",
            "participated": "Collaborated"
        }
        
        cliches = ["team player", "hard worker", "outside the box", "synergy", "detail-oriented"]
        suggestions = []
        rewrite = clean_bullet
        
        if first_word in weak_to_strong:
            strong_verb = weak_to_strong[first_word]
            suggestions.append(f"- Replace '{first_word}' with '{strong_verb}' for more impact.")
            rewrite = re.sub(rf'^{first_word}', strong_verb, rewrite, count=1, flags=re.IGNORECASE)
            
        for cliche in cliches:
            if cliche in rewrite.lower():
                suggestions.append(f"- Remove the buzzword '{cliche}'.")
                rewrite = re.sub(cliche, "effective collaboration", rewrite, flags=re.IGNORECASE)
        
        # Inject Target Role Relevance if keyword not in bullet
        used_key = ""
        if role_keywords:
            kw = random.choice(role_keywords).capitalize()
            if kw.lower() not in rewrite.lower():
                suggestions.append(f"- Tailor to the '{target_role}' role by mentioning '{kw}'.")
                used_key = f", utilizing {kw} techniques"
                
        if not has_metrics:
            impact_phrases = [
                " resulting in a 20% increase in workflow efficiency.",
                " successfully reducing operational constraints by 15%.",
                " delivering the solution 10% ahead of the scheduled deadline.",
                " leading to a measurable improvement in critical objectives."
            ]
            suffix = random.choice(impact_phrases)
            suggestions.append("- Add a quantifiable achievement/metric.")
            rewrite = rewrite.rstrip('.') + used_key + "," + suffix
        else:
            if used_key:
                rewrite = rewrite.rstrip('.') + used_key + "."
                
        suggestion_text = "\n  ".join(suggestions) if suggestions else "- Bullet is solid, but dynamically aligned to role."
        
        improved.append(f"*{clean_bullet}*\n  **Suggestions:**\n  {suggestion_text}\n\n  [REWRITE]: {rewrite}")

    # Ensure we return at least 4 if the resume had fewer
    while len(improved) < 4:
        kw1 = random.choice(role_keywords) if role_keywords else "industry-standard"
        kw2 = random.choice(role_keywords) if len(role_keywords) > 1 else "data"
        tgt = target_role if target_role else "Target Role"
        dummy = f"Implemented {kw1.capitalize()} solutions utilizing {kw2.capitalize()}, optimizing workflows for the {tgt} position, resulting in 15% efficiency gain."
        improved.append(f"*New Specific Bullet Suggested*\n  **Suggestions:**\n  - Addition precisely tailored to {tgt}\n\n  [REWRITE]: {dummy}")

    # Limit exactly to 5 if it exceeded somehow
    return improved[:5]


# -------------------------------
# Generate Fallback Bullets
# -------------------------------
def generate_fallback_bullets(text, target_role=""):
    generated = []
    base_role = target_role if target_role else "operations"
    kw = base_role.split()[0].capitalize()
    
    generated.append(f"Spearheaded major initiatives matching {base_role} requirements, resulting in a 20% increase in overall productivity.")
    generated.append(f"Optimized internal workflows utilizing {kw} methodologies, successfully reducing operational overhead by 15%.")
    generated.append(f"Collaborated extensively with cross-functional teams to deliver {base_role} projects 2 weeks ahead of schedule.")
    generated.append(f"Engineered standard processes aligned with the {base_role} ecosystem, boosting efficiency.")
    generated.append(f"Pioneered data-driven strategies for {base_role}, exceeding quarterly performance metrics.")
    return generated


# -------------------------------
# ATS Match calculation
# -------------------------------
def ats_match(resume_text, jd_text, target_role=""):
    full_target = f"{target_role} {jd_text}"
    if not full_target.strip():
        return 0

    def clean(text):
        text = text.lower()
        text = re.sub(f'[{re.escape(string.punctuation)}]', ' ', text)
        words = text.split()
        stop_words = {"the", "and", "a", "to", "of", "in", "for", "is", "with", "on", "as", "an", "by", "at", "this", "that", "it", "or", "from", "be", "are", "you", "your", "we", "our", "will", "can", "their"}
        return {w for w in words if w not in stop_words and len(w) > 2}

    resume_words = clean(resume_text)
    jd_words = clean(full_target)

    if not jd_words:
        return 0

    match_count = len(resume_words.intersection(jd_words))
    # Small boost if target role words are explicitly in the resume
    role_words = clean(target_role)
    role_match = len(resume_words.intersection(role_words))
    
    match_percentage = ((match_count + role_match) / (len(jd_words) + max(1, len(role_words)))) * 100
    # Add a slight multiplier to compensate for strict word matches
    return min(100, round(match_percentage * 1.5))


# -------------------------------
# Navigation State
# -------------------------------
if "page" not in st.session_state:
    st.session_state.page = "upload"

def set_bg_from_local(image_file, add_box=False, blur=False, font_color=None):
    # Robust path resolution for Streamlit Cloud deployment
    base_path = os.path.dirname(os.path.abspath(__file__))
    absolute_path = os.path.join(base_path, image_file)
    
    encoded_string = ""
    if os.path.exists(absolute_path):
        with open(absolute_path, "rb") as file:
            encoded_string = base64.b64encode(file.read()).decode()
    elif os.path.exists(image_file):
        with open(image_file, "rb") as file:
            encoded_string = base64.b64encode(file.read()).decode()

    # Always generate CSS so styles don't break if image is missing
    css = "<style>\n.stApp {\n"
    if encoded_string:
        css += f"""
        background-image: url(data:image/png;base64,{encoded_string});
        """
    else:
        css += """
        background: radial-gradient(circle, #2C5364 0%, #203A43 50%, #0F2027 100%);
        """
    
    css += """
        background-size: cover;
        background-position: center;
        background-attachment: fixed;
    }
    """
    
    if add_box:
        css += """
        .block-container {
            background-color: rgba(255, 255, 255, 0.85);
            border-radius: 15px;
            padding: 3rem !important;
            box-shadow: 0 8px 32px rgba(0,0,0,0.15);
        }
        """
    elif blur:
        css += """
        .block-container {
            backdrop-filter: blur(15px);
            -webkit-backdrop-filter: blur(15px);
            background-color: rgba(50, 80, 90, 0.7);
            border-radius: 20px;
            padding: 3rem !important;
            box-shadow: 0 8px 32px rgba(0,0,0,0.3);
            border: 1px solid rgba(255,255,255,0.1);
        }
        """
        
    if font_color:
        css += f"""
        .block-container p, .block-container h1, .block-container h2, .block-container h3, 
        .block-container h4, .block-container h5, .block-container h6, .block-container span, 
        .block-container label, .subtitle, .main-title {{
            color: {font_color} !important;
        }}
        [data-testid="stFileUploadDropzone"] * {{
            color: #ffffff !important;
        }}
        """
        
    css += "</style>"
    st.markdown(css, unsafe_allow_html=True)

# -------------------------------
# Sidebar Diagnostics
# -------------------------------
with st.sidebar:
    st.markdown("### 🛠️ Diagnostics")
    with st.expander("System Checks"):
        st.write("**RapidOCR**:", "✅ Ready" if RAPIDOCR_AVAILABLE else "❌ Missing")
        try:
            import cv2
            st.write("**OpenCV**:", "✅ Ready")
        except:
            st.write("**OpenCV**:", "❌ Missing libgl1")
        try:
            from pdfplumber.utils import find_poppler_binaries
            st.write("**Poppler**:", "✅ Ready")
        except:
            st.write("**Poppler**:", "❌ Missing poppler-utils")
    
    if st.button("Reset Session"):
        st.session_state.clear()
        st.rerun()

# ===================================
# PAGE 1 : Upload Resume
# ===================================
if st.session_state.page == "upload":

    set_bg_from_local("background.png", blur=True, font_color="#ffffff")

    st.markdown('<div class="main-title">📄 Resume Analyzer & Improver</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitle">Enter your target job role, then upload your resume for ATS analysis.</div>', unsafe_allow_html=True)
    
    st.markdown("### 1. Target Job Role")
    st.markdown("What job role are you applying for? (e.g., Python Developer, Data Scientist)")
    target_role = st.text_input("target_role", label_visibility="collapsed", placeholder="Enter job role...")
    
    st.markdown("### 2. Job Description (Optional)")
    st.markdown("Paste the Job Description (JD) here for a more precise match.")
    jd = st.text_area("jd", label_visibility="collapsed", height=150)

    st.markdown("### 3. Upload Resume")
    st.markdown("Upload your resume in PDF, DOCX, JPG, or PNG format.")
    file = st.file_uploader("Upload Resume", type=["pdf","docx","jpg","png"], label_visibility="collapsed")
    
    if file:
        st.markdown("<br>", unsafe_allow_html=True)
        # Verify if job role is provided, notify user if missed for better analysis
        if not target_role:
            st.warning("⚠️ For best results, please provide a Target Job Role.")
            
        if st.button("🔍 Analyze Resume", use_container_width=True):
            st.session_state.text = extract_text(file)
            st.session_state.target_role = target_role
            st.session_state.jd = jd
            st.session_state.page = "analysis"
            st.rerun()


# ===================================
# PAGE 2 : Analysis
# ===================================
elif st.session_state.page == "analysis":

    set_bg_from_local("background2.png", blur=True, font_color="#ffffff")

    text = st.session_state.text
    target_role = st.session_state.get('target_role', '')
    jd = st.session_state.get('jd', '')

    # -------------------------------
    # Analysis Summary
    # -------------------------------
    st.markdown(f'<div class="section-title">Analysis Executive Summary {f"for {target_role}" if target_role else ""}</div>', unsafe_allow_html=True)
    
    words = len(text.split())
    if not text.strip() or words < 10:
        st.error("🚨 **Error**: No text could be extracted from your resume. This usually happens if the PDF is scanned and the server lacks OCR dependencies, or if the file is corrupted.")
        if st.button("Back to Upload"):
            st.session_state.page = "upload"
            st.rerun()
        st.stop()
    elif words < 100:
        st.warning("⚠️ **Warning**: The extracted text is very short. Analysis might be inaccurate.")
    
    col1_s, col2_s = st.columns(2)
    with col1_s:
        st.write(f"**Word Count**: {words}")
        st.write(f"**Readability**: {'High' if words > 300 else 'Medium'}")
    with col2_s:
        detected = [k for k, v in analyze_resume(text).items() if v]
        st.write(f"**Sections Found**: {len(detected)} / 6")
        st.write(f"**Status**: {'Ready for ATS' if len(detected) >= 4 else 'Needs Improvement'}")

    if st.button("Upload Another Resume"):
        st.session_state.page = "upload"
        st.rerun()

    # -------------------------------
    # ATS Match Score
    # -------------------------------
    if target_role or jd:
        st.markdown(f'<div class="section-title">Target Job ATS Match Score</div>', unsafe_allow_html=True)
        ats_score = ats_match(text, jd, target_role)
        
        st.progress(ats_score)
        st.markdown(f"### ATS Match to **{target_role if target_role else 'Job Description'}**: **{ats_score}%**")
        
        if ats_score >= 70:
            st.success("Great job! Your resume has a strong match with the target job keywords.")
        elif ats_score >= 40:
            st.warning("Moderate match. Try incorporating more exact keywords related to the target job role.")
        else:
            st.error("Low match. Please tailor your resume keywords to this specific job role.")

    # -------------------------------
    # Resume Text
    # -------------------------------
    st.markdown('<div class="section-title">Extracted Resume Text</div>', unsafe_allow_html=True)
    
    escaped_text = html.escape(text)

    st.markdown(f"""
    <div class="text-container">
    {escaped_text}
    </div>
    """, unsafe_allow_html=True)


    # -------------------------------
    # Section Detection
    # -------------------------------
    st.markdown('<div class="section-title">Resume Sections</div>', unsafe_allow_html=True)

    sections = analyze_resume(text)

    for sec, present in sections.items():
        if present:
            st.success(f"{sec} section detected")
        else:
            st.warning(f"{sec} section missing")


    # -------------------------------
    # Resume Score
    # -------------------------------
    st.markdown('<div class="section-title">Resume Overview Score</div>', unsafe_allow_html=True)

    details, total = calculate_score(sections, text, target_role)

    col1, col2 = st.columns(2)
    with col1:
        st.markdown(f"""
        <div class="metric-card">
            <h4>Formatting</h4>
            <h2>{details['Formatting']} <span style="font-size:16px;color:#e0e0e0;">/ 20</span></h2>
        </div>
        """, unsafe_allow_html=True)
        st.markdown(f"""
        <div class="metric-card">
            <h4>Content Quality</h4>
            <h2>{details['Content']} <span style="font-size:16px;color:#e0e0e0;">/ 30</span></h2>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown(f"""
        <div class="metric-card">
            <h4>Skills Relevance</h4>
            <h2>{details['Skills']} <span style="font-size:16px;color:#e0e0e0;">/ 25</span></h2>
        </div>
        """, unsafe_allow_html=True)
        st.markdown(f"""
        <div class="metric-card">
            <h4>Experience Impact</h4>
            <h2>{details['Experience']} <span style="font-size:16px;color:#e0e0e0;">/ 25</span></h2>
        </div>
        """, unsafe_allow_html=True)

    st.progress(total)

    st.markdown(f"""
    <div class="score-container">
    Final Resume Score: {total} / 100
    </div>
    """, unsafe_allow_html=True)


    # -------------------------------
    # Suggestions
    # -------------------------------
    st.markdown('<div class="section-title">Improvement Suggestions</div>', unsafe_allow_html=True)

    suggestions = generate_suggestions(sections, text, target_role)

    if suggestions:
        for s in suggestions:
            st.info(s)
    else:
        st.success("Your resume structure looks strong!")


    # -------------------------------
    # Bullet Improvements
    # -------------------------------
    st.markdown('<div class="section-title">Target Role Driven Bullet Rewrite</div>', unsafe_allow_html=True)

    bullets = extract_bullets(text)

    if bullets:
        improved = improve_bullets(bullets, target_role, jd)
        for b in improved:
            st.markdown(f"- {b}")
    else:
        example_bullets = generate_fallback_bullets(text, target_role)
        for b in example_bullets:
            st.markdown(f"- {b}")
        improved = []
        
    # -------------------------------
    # Download Report
    # -------------------------------
    st.markdown("---")
    
    report = f"RESUME REPORT AND BULLET REWRITES FOR: {target_role if target_role else 'General Role'}\n"
    report += "="*40 + "\n\n"
    
    report += f"ATS Match Score: {ats_score}%\n" if target_role or jd else ""
    report += f"Overall Score: {total} / 100\n\n"
    
    report += "[ IMPROVED BULLETS ]\n"
    
    if bullets and improved:
        for b in improved:
            if "[REWRITE]:" in b:
                parts = b.split("[REWRITE]:")
                if len(parts) > 1:
                    report += f"- {parts[1].strip()}\n\n"
            elif "➔ **Great**:" in b:
                match = re.search(r'^\*(.*?)\*', b)
                if match:
                    report += f"- {match.group(1)}\n\n"
    else:
        for eb in example_bullets:
            report += f"- {eb}\n\n"
            
    # Generate PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=11)
    
    clean_report_for_pdf = report.replace('➔', '->').replace('💡', 'Tip:').replace('📥', '')
    
    # Simple fallback for encoding errors in basic PDF gen
    clean_report_for_pdf = clean_report_for_pdf.encode('latin-1', 'replace').decode('latin-1')
    
    pdf.multi_cell(w=0, h=6, txt=clean_report_for_pdf)
    pdf_bytes = pdf.output()
            
    st.download_button(
        label="📥 Download Resume Report (PDF)",
        data=bytes(pdf_bytes),
        file_name=f"Resume_Report_{target_role.replace(' ','_') if target_role else 'Analytics'}.pdf",
        mime="application/pdf"
    )
